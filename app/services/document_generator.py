"""Serviço para gerar documentos estruturados (Excel, Word) a partir de respostas da IA."""

import re
from io import BytesIO
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter


class DocumentGenerator:
    """Gerador de documentos estruturados."""

    @staticmethod
    def parse_markdown_to_structure(text: str) -> Dict:
        """
        Parseia texto Markdown em estrutura de dados.
        
        Args:
            text: Texto em Markdown ou texto simples
            
        Returns:
            Dict com estrutura parseada
        """
        structure = {
            "title": "",
            "sections": [],
            "tables": [],
            "lists": [],
        }
        
        lines = text.split("\n")
        current_section = None
        current_list = None
        default_section = {"title": "Conteúdo", "content": [], "subsections": []}
        
        for line in lines:
            original_line = line
            line = line.strip()
            
            # Título principal (# Título)
            if line.startswith("# "):
                structure["title"] = line[2:].strip()
            
            # Subtítulos (## Subtítulo)
            elif line.startswith("## "):
                if current_list:
                    if current_section:
                        current_section["content"].append({"type": "list", "items": current_list})
                    else:
                        default_section["content"].append({"type": "list", "items": current_list})
                    current_list = None
                if current_section:
                    structure["sections"].append(current_section)
                current_section = {
                    "title": line[3:].strip(),
                    "content": [],
                    "subsections": [],
                }
            
            # Sub-subtítulos (### Sub-subtítulo)
            elif line.startswith("### "):
                if current_list:
                    if current_section:
                        current_section["content"].append({"type": "list", "items": current_list})
                    else:
                        default_section["content"].append({"type": "list", "items": current_list})
                    current_list = None
                if current_section:
                    current_section["subsections"].append({
                        "title": line[4:].strip(),
                        "content": [],
                    })
            
            # Lista numerada (1. item)
            elif re.match(r"^\d+\.\s+", line):
                if not current_list:
                    current_list = []
                current_list.append(re.sub(r"^\d+\.\s+", "", line))
            
            # Lista com marcador (- item ou * item)
            elif line.startswith("- ") or line.startswith("* "):
                if not current_list:
                    current_list = []
                current_list.append(line[2:].strip())
            
            # Tabela (detecta linhas com |)
            elif "|" in line and line.count("|") >= 2:
                if current_list:
                    if current_section:
                        current_section["content"].append({"type": "list", "items": current_list})
                    else:
                        default_section["content"].append({"type": "list", "items": current_list})
                    current_list = None
                # Detectar tabelas Markdown
                cells = [cell.strip() for cell in line.split("|") if cell.strip()]
                if cells and not all(c.startswith("-") for c in cells):  # Não é separador
                    if not structure["tables"] or len(structure["tables"][-1]) == 0:
                        structure["tables"].append([])
                    structure["tables"][-1].append(cells)
            
            # Texto normal
            elif line:
                if current_list:
                    if current_section:
                        current_section["content"].append({"type": "list", "items": current_list})
                    else:
                        default_section["content"].append({"type": "list", "items": current_list})
                    current_list = None
                # Adicionar texto à seção atual ou à seção padrão
                if current_section:
                    current_section["content"].append({"type": "text", "content": line})
                else:
                    default_section["content"].append({"type": "text", "content": line})
        
        # Finalizar lista pendente
        if current_list:
            if current_section:
                current_section["content"].append({"type": "list", "items": current_list})
            else:
                default_section["content"].append({"type": "list", "items": current_list})
        
        # Adicionar seção atual
        if current_section:
            structure["sections"].append(current_section)
        
        # Se não houver seções mas houver conteúdo na seção padrão, adicionar
        if not structure["sections"] and default_section["content"]:
            structure["sections"].append(default_section)
        
        return structure

    @staticmethod
    def generate_excel(text: str, filename: str = "documento") -> BytesIO:
        """
        Gera arquivo Excel a partir de texto estruturado.
        
        Args:
            text: Texto em Markdown ou texto simples
            filename: Nome do arquivo (sem extensão)
            
        Returns:
            BytesIO com o arquivo Excel
        """
        wb = Workbook()
        ws = wb.active
        ws.title = "Documento"
        
        # Estilos
        title_font = Font(bold=True, size=14, color="FFFFFF")
        title_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        header_font = Font(bold=True, size=12)
        header_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
        border = Border(
            left=Side(style="thin"),
            right=Side(style="thin"),
            top=Side(style="thin"),
            bottom=Side(style="thin"),
        )
        
        row = 1
        
        # Parsear estrutura
        structure = DocumentGenerator.parse_markdown_to_structure(text)
        
        # Título
        if structure["title"]:
            ws.merge_cells(f"A{row}:D{row}")
            cell = ws[f"A{row}"]
            cell.value = structure["title"]
            cell.font = title_font
            cell.fill = title_fill
            cell.alignment = Alignment(horizontal="center", vertical="center")
            row += 2
        elif text.strip():  # Se não houver título mas houver texto, usar primeira linha como título
            first_line = text.strip().split("\n")[0][:100]  # Limitar a 100 caracteres
            ws.merge_cells(f"A{row}:D{row}")
            cell = ws[f"A{row}"]
            cell.value = first_line
            cell.font = title_font
            cell.fill = title_fill
            cell.alignment = Alignment(horizontal="center", vertical="center")
            row += 2
        
        # Seções
        for section in structure["sections"]:
            # Título da seção
            if section["title"]:
                ws.merge_cells(f"A{row}:D{row}")
                cell = ws[f"A{row}"]
                cell.value = section["title"]
                cell.font = header_font
                cell.fill = header_fill
                row += 1
            
            # Conteúdo da seção
            for content_item in section["content"]:
                if content_item["type"] == "text":
                    ws[f"A{row}"] = content_item["content"]
                    row += 1
                elif content_item["type"] == "list":
                    for item in content_item["items"]:
                        ws[f"B{row}"] = f"• {item}"
                        row += 1
                    # Espaçamento após lista
                    row += 1
            
            # Sub-seções
            for subsection in section.get("subsections", []):
                if subsection["title"]:
                    ws[f"B{row}"] = subsection["title"]
                    cell = ws[f"B{row}"]
                    cell.font = Font(bold=True, size=11)
                    row += 1
                for content_item in subsection.get("content", []):
                    if content_item["type"] == "text":
                        ws[f"C{row}"] = content_item["content"]
                        row += 1
                    elif content_item["type"] == "list":
                        for item in content_item["items"]:
                            ws[f"C{row}"] = f"  • {item}"
                            row += 1
            
            row += 1
        
        # Tabelas
        for table in structure["tables"]:
            if not table:
                continue
            
            start_row = row
            for i, row_data in enumerate(table):
                for j, cell_data in enumerate(row_data):
                    cell = ws.cell(row=row, column=j + 1, value=cell_data)
                    cell.border = border
                    if i == 0:  # Cabeçalho
                        cell.font = header_font
                        cell.fill = header_fill
                    ws.column_dimensions[get_column_letter(j + 1)].width = 20
                row += 1
            row += 2
        
        # Ajustar largura das colunas
        ws.column_dimensions["A"].width = 30
        ws.column_dimensions["B"].width = 30
        ws.column_dimensions["C"].width = 30
        ws.column_dimensions["D"].width = 30
        
        # Salvar em BytesIO
        output = BytesIO()
        wb.save(output)
        output.seek(0)
        return output

    @staticmethod
    def _add_formatted_text_to_paragraph(paragraph, text: str):
        """
        Adiciona texto formatado a um parágrafo existente, convertendo Markdown.
        
        Args:
            paragraph: Parágrafo do documento Word
            text: Texto com Markdown
        """
        import re
        
        # Remover emojis
        text = re.sub(r'[📋🌐🏛️🏥📝📄💙👋🩺✅❌⚠️💡]', '', text)
        
        # Se não tiver formatação, adicionar como texto simples
        if '**' not in text and '*' not in text and '`' not in text:
            paragraph.add_run(text)
            return
        
        # Processar formatação: negrito primeiro, depois itálico
        # Dividir por negrito (**texto**)
        parts = re.split(r'(\*\*[^*]+\*\*)', text)
        
        for part in parts:
            if not part:
                continue
            
            if part.startswith('**') and part.endswith('**'):
                # Negrito
                inner_text = part[2:-2]
                # Verificar se tem itálico dentro
                if '*' in inner_text:
                    italic_parts = re.split(r'(\*[^*]+\*)', inner_text)
                    for ip in italic_parts:
                        if ip.startswith('*') and ip.endswith('*'):
                            run = paragraph.add_run(ip[1:-1])
                            run.bold = True
                            run.italic = True
                        elif ip:
                            run = paragraph.add_run(ip)
                            run.bold = True
                else:
                    run = paragraph.add_run(inner_text)
                    run.bold = True
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                # Itálico (só se não for negrito)
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif '`' in part:
                # Código inline
                code_parts = re.split(r'(`[^`]+`)', part)
                for cp in code_parts:
                    if cp.startswith('`') and cp.endswith('`'):
                        run = paragraph.add_run(cp[1:-1])
                        run.font.name = 'Courier New'
                    elif cp:
                        paragraph.add_run(cp)
            else:
                # Texto normal
                paragraph.add_run(part)

    @staticmethod
    def generate_word(text: str, filename: str = "documento") -> BytesIO:
        """
        Gera arquivo Word a partir de texto estruturado, convertendo Markdown.
        
        Args:
            text: Texto em Markdown ou texto simples
            filename: Nome do arquivo (sem extensão)
            
        Returns:
            BytesIO com o arquivo Word
        """
        import re
        from docx.shared import Pt, RGBColor
        
        doc = Document()
        
        # Processar texto linha por linha
        lines = text.split('\n')
        
        for i, line in enumerate(lines):
            original_line = line
            line = line.strip()
            
            if not line:
                # Linha vazia - adicionar espaço
                doc.add_paragraph()
                continue
            
            # Remover emojis
            line_clean = re.sub(r'[📋🌐🏛️🏥📝📄💙👋🩺✅❌⚠️💡]', '', line)
            
            # Título principal (# Título)
            if line.startswith('# ') and not line.startswith('##'):
                title_text = line_clean[2:].strip()
                # Remover formatação do título
                title_text = re.sub(r'\*\*|\*|`', '', title_text)
                title = doc.add_heading(title_text, 0)
                title.alignment = 1  # Centralizado
            
            # Subtítulo (## Subtítulo)
            elif line.startswith('## ') and not line.startswith('###'):
                subtitle_text = line_clean[3:].strip()
                # Remover formatação
                subtitle_text = re.sub(r'\*\*|\*|`', '', subtitle_text)
                doc.add_heading(subtitle_text, 1)
            
            # Sub-subtítulo (### Subtítulo)
            elif line.startswith('### '):
                subsubtitle_text = line_clean[4:].strip()
                # Remover formatação
                subsubtitle_text = re.sub(r'\*\*|\*|`', '', subsubtitle_text)
                doc.add_heading(subsubtitle_text, 2)
            
            # Lista numerada (1. item)
            elif re.match(r'^\d+\.\s+', line):
                item_text = re.sub(r'^\d+\.\s+', '', line_clean)
                # Criar parágrafo com estilo de lista
                p = doc.add_paragraph()
                DocumentGenerator._add_formatted_text_to_paragraph(p, item_text)
                p.style = "List Number"
            
            # Lista com marcador (- item ou * item, mas não **negrito**)
            elif (line.startswith('- ') or (line.startswith('* ') and not line.startswith('**'))):
                item_text = line_clean[2:].strip()
                # Criar parágrafo com estilo de lista
                p = doc.add_paragraph()
                DocumentGenerator._add_formatted_text_to_paragraph(p, item_text)
                p.style = "List Bullet"
            
            # Lista com indentação (4 espaços ou tab)
            elif line.startswith('    ') or line.startswith('\t'):
                item_text = line_clean.strip()
                # Remover marcador de lista se houver
                if item_text.startswith('• ') or item_text.startswith('- '):
                    item_text = item_text[2:].strip()
                p = doc.add_paragraph()
                DocumentGenerator._add_formatted_text_to_paragraph(p, item_text)
                p.style = "List Bullet"
                # Adicionar indentação
                p.paragraph_format.left_indent = Inches(0.5)
            
            # Texto normal com formatação
            else:
                # Verificar se é uma linha especial (ex: "**PRESCRIÇÃO:**" ou "**PRESCRIÇÃO PARA...**")
                if '**' in line_clean:
                    # Verificar se é um título em negrito (começa e termina com ** ou tem ** no início)
                    if line_clean.startswith('**') and (line_clean.endswith('**') or ':**' in line_clean):
                        # É um título em negrito
                        title_text = re.sub(r'\*\*|:', '', line_clean).strip()
                        heading = doc.add_heading(title_text, 2)
                        for run in heading.runs:
                            run.bold = True
                    else:
                        # Parágrafo normal com formatação
                        p = doc.add_paragraph()
                        DocumentGenerator._add_formatted_text_to_paragraph(p, line_clean)
                else:
                    # Parágrafo normal sem formatação Markdown
                    p = doc.add_paragraph(line_clean)
        
        # Salvar em BytesIO
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output

    @staticmethod
    def generate_from_ai_response(
        ai_response: str,
        format: str = "excel",
        filename: str = "resposta_ia",
    ) -> BytesIO:
        """
        Gera documento a partir de resposta da IA.
        
        Args:
            ai_response: Resposta da IA (texto ou Markdown)
            format: Formato do documento ("excel" ou "word")
            filename: Nome do arquivo (sem extensão)
            
        Returns:
            BytesIO com o arquivo gerado
        """
        if format.lower() == "excel" or format.lower() == "xlsx":
            return DocumentGenerator.generate_excel(ai_response, filename)
        elif format.lower() == "word" or format.lower() == "docx":
            return DocumentGenerator.generate_word(ai_response, filename)
        else:
            raise ValueError(f"Formato não suportado: {format}. Use 'excel' ou 'word'.")

